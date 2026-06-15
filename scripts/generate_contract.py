#!/usr/bin/env python3
"""
generate_contract.py — 从飞书合同信息收集表读取记录，匹配合同模板，填充变量，
插入证件扫描件，生成 docx，可选发送邮件。

用法：
    python3 scripts/generate_contract.py --list
    python3 scripts/generate_contract.py --name "测试候选人B"
    python3 scripts/generate_contract.py --name "测试候选人B" --dry-run
    python3 scripts/generate_contract.py --name "测试候选人B" --send
    python3 scripts/generate_contract.py --name "测试候选人B" --draft
"""

import sys, re, json, io, os, argparse, subprocess, tempfile
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image

sys.path.insert(0, str(Path(__file__).parent))
from config_loader import load_config, get_smtp, get_lark, get_paths, is_test_mode, get_test_email
from field_resolver import field_id_or
from lark_cli_utils import normalize_record_list_data, run_lark_cli_json
from manual_trace import log_manual_step
from field_mapping import (
    FORM_BASE_TOKEN, FORM_TABLE_ID,
    ACCOUNT_TYPE_FIELD_ID,
    COMMON_FORM_FIELDS, ATTACHMENT_FIELDS,
    VM_INPUT_VARS, AUTO_VARS,
    get_all_form_fields, is_company_contract,
)

_CFG = load_config()

# ── 配置（全部从 config.yaml 读取，不再硬编码）──────────────────
_lark = get_lark(_CFG)

COLLECT_BASE   = _lark.get("contract_base_token") or _lark.get("base_token", "")
COLLECT_TABLE  = _lark.get("contract_table_id", "")

TEMPLATE_BASE  = _lark.get("template_base_token", "")
TEMPLATE_TABLE = _lark.get("template_table_id", "")
TEMPLATE_ATT_FLD = "fldSeCxvVQ"                    # AI合同模版 附件字段
TEMPLATE_NAME_FLD = "fldppDlNai"                   # 合同名1（公式字段）
TEMPLATE_VARS_FLD = "fldPx6ZMh8"                   # 所需变量（文本字段）

OUTPUT_DIR = Path(get_paths(_CFG).get("contract_output", "~/Documents/loc-contracts/output/")).expanduser()

TEST_MODE  = is_test_mode(_CFG)
TEST_EMAIL = get_test_email(_CFG)

FLD_NAME      = field_id_or("contract_info", "contract.name", "fld2JEyq9H")
FLD_ACCT_NAME = field_id_or("contract_info", "contract.bank_account_name", "fldvZMzuk3")
FLD_EMAIL     = field_id_or("contract_info", "contract.email", "fldYELKkKa")
FLD_SIGNED    = field_id_or("contract_info", "contract.signed_contract", "fldj4zCL5L")
FLD_ID_SCAN   = field_id_or("contract_info", "contract.id_scan", "fldia8GcRh")
FLD_ID_NO     = field_id_or("contract_info", "contract.id_number", "fld3hdHuVd")
FLD_BANK_NAME = field_id_or("contract_info", "contract.bank_name", "fldyPyrLdp")
FLD_BANK_ADDR = field_id_or("contract_info", "contract.bank_address", "fldDLk0Jh9")
FLD_CURRENCY  = field_id_or("contract_info", "contract.currency", "fldSZE1Shy")


# ── lark-cli 工具 ──────────────────────────────────────────────
def lark(*args) -> dict:
    resp = run_lark_cli_json(*args)
    if not isinstance(resp, dict):
        raise RuntimeError(f"lark-cli 非 JSON:\n{str(resp)[:300]}")
    return resp


def lark_download_attachment(
    file_token: str,
    dest: Path,
    *,
    base_token: str,
    table_id: str,
    record_id: str,
):
    """下载飞书 Base 附件到本地路径。"""
    r = subprocess.run(
        ["lark-cli", "base", "+record-download-attachment",
         "--base-token", base_token,
         "--table-id", table_id,
         "--record-id", record_id,
         "--file-token", file_token,
         "--output", dest.name],
        capture_output=True, text=True,
        cwd=str(dest.parent),
    )
    if r.returncode != 0:
        raise RuntimeError(f"附件下载失败: {r.stderr.strip()}")


def find_table_id_by_name(base_token: str, table_name: str) -> str:
    """按表名查找 table_id，用于 config.yaml 中旧 table_id 失效时兜底。"""
    try:
        resp = lark("base", "+table-list", "--base-token", base_token, "--format", "json")
    except Exception:
        return ""
    data = resp.get("data", resp)
    raw_tables = []
    for key in ("items", "tables", "table_list"):
        if isinstance(data, dict) and isinstance(data.get(key), list):
            raw_tables = data[key]
            break
    if not raw_tables and isinstance(resp.get("items"), list):
        raw_tables = resp["items"]
    for table in raw_tables:
        name = table.get("name") or table.get("table_name") or ""
        table_id = table.get("table_id") or table.get("id") or ""
        if name == table_name and table_id:
            return table_id
    return ""


# ── 记录拉取 ──────────────────────────────────────────────────
def fetch_collect_records() -> list:
    global COLLECT_TABLE
    records, page_token = [], None
    while True:
        args = ["base", "+record-list",
                "--base-token", COLLECT_BASE,
                "--table-id", COLLECT_TABLE,
                "--format", "json", "--limit", "100"]
        if page_token:
            args += ["--page-token", page_token]
        try:
            resp = lark(*args)
        except RuntimeError as e:
            fallback_table = find_table_id_by_name(COLLECT_BASE, "合同信息收集")
            if fallback_table and fallback_table != COLLECT_TABLE:
                print(f"⚠️  合同信息表 ID 失效，已按表名 fallback：{fallback_table}")
                COLLECT_TABLE = fallback_table
                records, page_token = [], None
                continue
            raise e
        db = resp["data"]
        records.extend(normalize_record_list_data(db))
        if not db.get("has_more") or not db.get("page_token"):
            break
        page_token = db["page_token"]
    return records


def fetch_template_records() -> list:
    resp = lark("base", "+record-list",
                "--base-token", TEMPLATE_BASE,
                "--table-id", TEMPLATE_TABLE,
                "--format", "json")
    records = []
    data = resp.get("data", {})
    if data.get("data"):
        records.extend(normalize_record_list_data(data))
        return records
    for rec in resp.get("data", {}).get("records", []):
        records.append(rec)
    return records


# ── 字段值提取 ────────────────────────────────────────────────
def extract_text(val) -> str:
    if not val:
        return ""
    if isinstance(val, list):
        parts = []
        for v in val:
            if isinstance(v, dict):
                parts.append(v.get("name") or v.get("text") or "")
            else:
                parts.append(str(v))
        text = " ".join(p for p in parts if p).strip()
    else:
        text = str(val).strip()
    text = re.sub(r'\[([^\]]+)\]\([^)]*\)', r'\1', text)
    return text


def is_china_id_number(text: str) -> bool:
    """Return True for common mainland China resident ID format."""
    value = extract_text(text).replace(" ", "")
    return bool(re.fullmatch(r"\d{17}[\dXx]", value))


def has_domestic_bank_signal(text: str) -> bool:
    value = extract_text(text).lower()
    markers = [
        "中国", "china", "beijing", "shanghai", "guangzhou", "shenzhen",
        "hangzhou", "nanjing", "chengdu", "wuhan", "xiamen", "anhui", "hefei",
        "工商银行", "农业银行", "中国银行", "建设银行", "招商银行", "交通银行",
        "邮储", "中信", "光大", "浦发", "兴业", "民生", "广发", "平安银行",
    ]
    return any(marker in value for marker in markers)


def is_domestic_personal_account(fields: dict) -> bool:
    acct_type = extract_text(fields.get(ACCOUNT_TYPE_FIELD_ID, ""))
    if not ("个人" in acct_type or "personal" in acct_type.lower()):
        return False
    id_no = extract_text(fields.get(FLD_ID_NO, ""))
    bank_name = extract_text(fields.get(FLD_BANK_NAME, ""))
    bank_addr = extract_text(fields.get(FLD_BANK_ADDR, ""))
    return (
        is_china_id_number(id_no) or
        has_domestic_bank_signal(bank_name) or
        has_domestic_bank_signal(bank_addr)
    )


def is_personal_account(fields: dict) -> bool:
    acct_type = extract_text(fields.get(ACCOUNT_TYPE_FIELD_ID, ""))
    return "个人" in acct_type or "personal" in acct_type.lower()


def is_company_account(fields: dict) -> bool:
    acct_type = extract_text(fields.get(ACCOUNT_TYPE_FIELD_ID, ""))
    return "公司" in acct_type or "business" in acct_type.lower()


def is_cny_currency(fields: dict) -> bool:
    currency = extract_text(fields.get(FLD_CURRENCY, "")).lower()
    return "人民币" in currency or "cny" in currency or "rmb" in currency


def is_foreign_currency(fields: dict) -> bool:
    currency = extract_text(fields.get(FLD_CURRENCY, "")).lower()
    if not currency:
        return False
    foreign_markers = [
        "外币", "usd", "美元", "eur", "欧元", "jpy", "日元", "krw", "韩元",
        "hkd", "港币", "gbp", "英镑", "sgd", "新币", "aud", "cad",
    ]
    return any(marker in currency for marker in foreign_markers) and not is_cny_currency(fields)


def score_contract_template(name: str, fields: dict) -> int:
    """Score how well a template name matches account type, region, and currency.

    Region and currency are intentionally separate. A domestic account can receive
    foreign currency, and an overseas account can receive CNY.
    """
    lowered = name.lower()
    is_company_acct = is_company_account(fields)
    is_personal_acct = is_personal_account(fields)
    is_domestic_personal = is_domestic_personal_account(fields)
    wants_cny = is_cny_currency(fields)
    wants_foreign = is_foreign_currency(fields)

    score = 0

    if is_company_acct:
        if "公司" in name or "business" in lowered or "company" in lowered:
            score += 40
        if "个人" in name or "personal" in lowered:
            score -= 20
        return score

    if is_personal_acct:
        if "个人" in name or "personal" in lowered:
            score += 20
        if "公司" in name or "business" in lowered or "company" in lowered:
            score -= 30

    if is_domestic_personal:
        if "境内个人" in name:
            score += 30
        if "境外" in name or "foreign" in lowered:
            score -= 20
    else:
        if "境外" in name or "foreign" in lowered:
            score += 20
        if "境内个人" in name:
            score -= 15

    if wants_cny:
        if "人民币" in name or "cny" in lowered or "rmb" in lowered:
            score += 30
        if "外币" in name or "foreign currency" in lowered:
            score -= 25
    elif wants_foreign:
        if "外币" in name or "foreign currency" in lowered:
            score += 30
        if "人民币" in name or "cny" in lowered or "rmb" in lowered:
            score -= 25

    return score


def extract_attachments(val) -> list:
    """返回附件列表 [{name, file_token, size}]"""
    if not val or not isinstance(val, list):
        return []
    result = []
    for item in val:
        if isinstance(item, dict) and item.get("file_token"):
            result.append(item)
    return result


def open_docx(path: Path) -> bool:
    """Open generated docx, preferring WPS when installed on macOS."""
    candidates = []
    wps_path = Path("/Applications/wpsoffice.app")
    if wps_path.exists():
        candidates.append(["open", "-a", str(wps_path), str(path)])
    candidates.extend([
        ["open", "-b", "com.kingsoft.wpsoffice.mac", str(path)],
        ["open", str(path)],
    ])
    last_error = ""
    for cmd in candidates:
        result = subprocess.run(cmd, capture_output=True, text=True)
        if result.returncode == 0:
            if "wpsoffice" in " ".join(cmd).lower():
                print("✅ 已使用 WPS 打开合同预览")
            else:
                print("✅ 已打开合同预览")
            return True
        last_error = (result.stderr or result.stdout or "").strip()
    print(f"⚠️  自动打开失败，请手动打开：{path}")
    if last_error:
        print(f"   系统返回：{last_error}")
    return False


# ── 模板匹配 ──────────────────────────────────────────────────
def match_template(template_records: list, contract_name: str) -> dict | None:
    """按合同名精确匹配模板记录"""
    for rec in template_records:
        name1 = extract_text(rec.get("fields", {}).get(TEMPLATE_NAME_FLD, ""))
        if name1 and name1.replace(".docx", "") == contract_name.replace(".docx", ""):
            return rec
    return None


def pick_template_for_candidate(template_records: list, fields: dict, auto_confirm: bool = False) -> tuple[dict | None, str]:
    """
    根据候选人账户类型自动推荐合同模板，VM 可直接回车确认或手动选择。
    """
    available = []
    for rec in template_records:
        name = extract_text(rec.get("fields", {}).get(TEMPLATE_NAME_FLD, ""))
        agt_raw = rec.get("fields", {}).get("fldS1Wlc2x", "")
        agt_id  = extract_text(agt_raw) if isinstance(agt_raw, (str, list)) else ""
        if not name or "AGT-009" in str(agt_id):
            continue
        available.append((rec, name))

    if not available:
        print("❌ 合同模板表中无可用模板")
        return None, ""

    available.sort(key=lambda x: score_contract_template(x[1], fields), reverse=True)

    print("\n可用合同模板：")
    for i, (rec, name) in enumerate(available, 1):
        tag = "  ← 推荐" if i == 1 else ""
        print(f"  {i:>2}. {name}{tag}")

    if auto_confirm:
        print("\n--yes 模式：自动使用推荐模板 [1]")
        choice = "1"
    else:
        choice = input(f"\n请输入模板序号（直接回车使用推荐 [1]）: ").strip()
    if choice == "":
        choice = "1"
    try:
        idx = int(choice) - 1
        rec, name = available[idx]
        return rec, name
    except (ValueError, IndexError):
        print("❌ 输入无效，已取消")
        return None, ""


# ── 变量构建 ──────────────────────────────────────────────────
def build_var_map(fields: dict, required_vars: list, vm_overrides: dict = None) -> tuple[dict, list]:
    """
    根据收集表字段 + 模板所需变量列表，构建 {{{变量名}}: 值} 替换字典。
    required_vars: 从模板所需变量字段解析出的列表，如 ["乙方姓名", "合同生效日期", ...]
    返回: (var_map, missing_list)
    """
    vm_overrides = vm_overrides or {}
    contract_defaults = _CFG.get("contract_defaults", {}) or {}

    # 获取账户类型，决定银行字段路由
    acct_type_raw = extract_text(fields.get(ACCOUNT_TYPE_FIELD_ID, ""))
    try:
        form_fields = get_all_form_fields(acct_type_raw)
    except ValueError:
        # 账户类型为空时默认走个人账户
        from field_mapping import PERSONAL_ACCOUNT_FIELDS
        form_fields = {**COMMON_FORM_FIELDS, **PERSONAL_ACCOUNT_FIELDS}

    var_map   = {}   # {"{{变量名}}": "值"}
    filled    = []   # 成功填充的变量名
    empty     = []   # 字段存在但值为空（资源商未填）
    unmatched = []   # 变量名在映射表中找不到对应字段

    # 处理签署日期自动拆分
    sign_date = vm_overrides.get("签署日期", datetime.now().strftime("%Y-%m-%d"))
    try:
        sd = datetime.strptime(sign_date, "%Y-%m-%d")
        auto_vals = {
            "签署年": str(sd.year),
            "签署月": f"{sd.month:02d}",
            "签署日": f"{sd.day:02d}",
            "甲方合同编号": "",
        }
    except ValueError:
        auto_vals = {"签署年": "", "签署月": "", "签署日": "", "甲方合同编号": ""}

    for var in required_vars:
        key = f"{{{{{var}}}}}"   # "{{变量名}}"

        # 1. VM 手动覆盖
        if var in vm_overrides:
            var_map[key] = str(vm_overrides[var])
            filled.append(var)
            continue

        # 2. 自动计算
        if var in auto_vals:
            var_map[key] = auto_vals[var]
            filled.append(var)
            continue

        # 2.5 本机固定合同变量
        if var in contract_defaults:
            var_map[key] = str(contract_defaults[var])
            filled.append(var)
            continue

        # 3. 收集表字段
        if var in form_fields:
            val = extract_text(fields.get(form_fields[var], ""))
            var_map[key] = val
            if val:
                filled.append(var)
            else:
                empty.append(var)   # 字段有映射，但资源商未填值
            continue

        # 4. 乙方签署 = 乙方姓名
        if var == "乙方签署":
            val = extract_text(fields.get(COMMON_FORM_FIELDS["乙方姓名"], ""))
            var_map[key] = val
            (filled if val else empty).append(var)
            continue

        # 5. 映射表中找不到该变量 → unmatched
        var_map[key] = ""
        unmatched.append(var)

    missing = empty + unmatched   # 兼容旧调用方
    return var_map, missing, filled, empty, unmatched


def parse_required_vars(vars_text: str) -> list:
    """从「所需变量」字段文本解析变量名列表"""
    # 格式: {{乙方姓名}} / {{合同生效日期}} / ...
    return re.findall(r'\{\{([^}]+)\}\}', vars_text)


# ── docx 文本替换 ─────────────────────────────────────────────
def replace_para(para, var_map: dict):
    full = "".join(r.text for r in para.runs)
    if not any(var in full for var in var_map):
        return
    merged = full
    for var, val in var_map.items():
        merged = merged.replace(var, val)
    if para.runs:
        para.runs[0].text = merged
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.text = merged


def fill_template_vars(doc: Document, var_map: dict):
    for para in doc.paragraphs:
        replace_para(para, var_map)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_para(para, var_map)


# ── 证件扫描件插入 ────────────────────────────────────────────
def insert_id_scan(doc: Document, image_paths: list[Path]):
    """
    找到合同末尾的附件描述段落（身份证/护照相关），在其后插入图片。
    支持多张图（正反面）。
    """
    from field_mapping import ATTACHMENT_FIELDS
    keywords = ATTACHMENT_FIELDS["乙方证件扫描件"]["anchor_keywords"]

    # 找到锚点段落的索引
    anchor_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if any(kw in text for kw in keywords):
            anchor_idx = i

    if anchor_idx is None:
        print("  ⚠️  未找到证件附件锚点段落，图片将追加到文档末尾")
        anchor_idx = len(doc.paragraphs) - 1

    # 在锚点段落后插入图片
    anchor_para = doc.paragraphs[anchor_idx]
    for img_path in image_paths:
        # 检查图片尺寸，超宽则限制宽度
        try:
            with Image.open(img_path) as im:
                w, h = im.size
                # A4 可用宽约 16cm，限制在 15cm 内
                max_w_inches = 5.9
                img_w_inches = min(w / 96, max_w_inches)  # 假设 96dpi
        except Exception:
            img_w_inches = 5.5

        # python-docx 插入图片需要在段落后新增段落
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy

        new_para = OxmlElement('w:p')
        anchor_para._element.addnext(new_para)
        # 重新获取刚插入的段落对象
        inserted = None
        for j, p in enumerate(doc.paragraphs):
            if p._element is new_para:
                inserted = p
                break

        if inserted is None:
            # fallback: 追加
            inserted = doc.add_paragraph()

        run = inserted.add_run()
        run.add_picture(str(img_path), width=Inches(img_w_inches))
        print(f"  ✅ 插入图片：{img_path.name}（宽 {img_w_inches:.1f} 英寸）")


def download_id_scan_images(fields: dict, tmpdir: Path, record_id: str) -> list[Path]:
    """从飞书下载证件扫描件图片到临时目录，返回本地路径列表"""
    attachments = extract_attachments(fields.get(FLD_ID_SCAN, []))
    if not attachments:
        return []

    paths = []
    for att in attachments:
        file_token = att.get("file_token", "")
        filename   = att.get("name", f"id_scan_{file_token[:8]}.jpg")
        dest = tmpdir / filename
        try:
            lark_download_attachment(
                file_token,
                dest,
                base_token=COLLECT_BASE,
                table_id=COLLECT_TABLE,
                record_id=record_id,
            )
            paths.append(dest)
            print(f"  ✅ 下载证件扫描件：{filename}")
        except Exception as e:
            print(f"  ⚠️  下载失败（{filename}）：{e}")
    return paths


# ── 银行账户名校验 ────────────────────────────────────────────
def check_name_match(full_name: str, bank_name: str) -> tuple[bool, str]:
    if not full_name or not bank_name:
        return True, "⚠️  姓名或账户名为空，请人工确认"
    if full_name.strip().lower() == bank_name.strip().lower():
        return False, "✅ 姓名与账户名一致"
    has_chinese = bool(re.search(r'[\u4e00-\u9fff]', full_name))
    if has_chinese:
        return True, f"⚠️  中文姓名「{full_name}」对应账户名「{bank_name}」，请确认拼音正确"
    name_parts = set(full_name.upper().split())
    bank_parts = set(bank_name.upper().split())
    if name_parts == bank_parts:
        return False, "✅ 姓名与账户名一致（词序不同）"
    if name_parts & bank_parts:
        return True, f"⚠️  姓名「{full_name}」与账户名「{bank_name}」部分匹配，请确认"
    return True, f"❌ 姓名「{full_name}」与账户名「{bank_name}」不匹配，请检查"


# ── 邮件发送 ──────────────────────────────────────────────────
def send_email(to_email: str, name: str, contract_path: Path, lang: str = "zh", draft: bool = False):
    import smtplib, ssl
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders

    smtp = get_smtp(_CFG)
    actual_to = TEST_EMAIL if TEST_MODE else to_email
    if not draft and not TEST_MODE:
        raise RuntimeError("生产环境禁止直接发送合同邮件；请使用 --draft 生成草稿，由 VM 人工检查后发送。")

    if lang == "zh":
        subject = f"【37GAMES】翻译委托框架协议 - {name}"
        body = (
            f"您好，\n\n感谢您提交合同信息。\n\n"
            f"请查收附件中的翻译委托框架协议，合同中的个人信息已预先填写。\n\n"
            f"请您：\n"
            f"1. 核对合同中的个人信息是否正确\n"
            f"2. 打印合同\n"
            f"3. 在乙方签名栏亲笔签字\n"
            f"4. 填写合同生效日期（即签字当天日期）\n"
            f"5. 扫描或拍照已签字页，发回此邮箱\n\n"
            f"如有任何问题，欢迎随时联系。\n\n"
            f"37GAMES 本地化团队"
        )
    else:
        subject = f"[37GAMES] Service Agreement - {name}"
        body = (
            f"Dear {name},\n\n"
            f"Thank you for submitting your contract information.\n\n"
            f"Please find the attached agreement with your personal details pre-filled.\n\n"
            f"Please:\n"
            f"1. Review and verify your information\n"
            f"2. Print the contract\n"
            f"3. Sign in the designated field\n"
            f"4. Fill in the effective date (today's date)\n"
            f"5. Scan the signed page and send it back\n\n"
            f"Best regards,\n37GAMES Localization Team"
        )

    msg = MIMEMultipart()
    msg["From"]    = smtp.get('user', '')
    msg["To"]      = actual_to
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "plain", "utf-8"))

    with open(contract_path, "rb") as fh:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(fh.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f'attachment; filename="{contract_path.name}"')
    msg.attach(part)

    if draft:
        import re as _re
        draft_dir = Path(get_paths(_CFG).get("contract_output", "~/Documents/loc-contracts/output/")).expanduser() / "drafts"
        draft_dir.mkdir(parents=True, exist_ok=True)
        safe = _re.sub(r'[^\w\-\u4e00-\u9fff.]', '_', name)
        draft_path = draft_dir / f"合同_{safe}.eml"
        draft_path.write_text(msg.as_string(), encoding="utf-8")
        print(f"📝 草稿已保存：{draft_path}")
        print(f"   双击用邮件客户端打开，确认无误后点发送")
        return

    ctx = ssl.SSLContext(ssl.PROTOCOL_TLS_CLIENT)
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    with smtplib.SMTP_SSL(smtp["host"], smtp.get("port", 465), context=ctx) as s:
        s.login(smtp["user"], smtp["password"])
        s.sendmail(smtp["user"], actual_to, msg.as_string())

    if TEST_MODE:
        print(f"⚠️  [测试模式] 邮件已发到 {actual_to}（原始目标：{to_email}）")
    else:
        print(f"✅ 邮件已发送至 {actual_to}")


# ── 列表展示 ──────────────────────────────────────────────────
def list_records(records):
    print(f"{'#':<4} {'record_id':<22} {'姓名':<20} {'邮箱':<32} {'证件扫描件'}")
    print("-" * 100)
    for i, rec in enumerate(records, 1):
        f = rec["fields"]
        name  = extract_text(f.get(FLD_NAME, ""))
        email = extract_text(f.get(FLD_EMAIL, ""))
        scans = extract_attachments(f.get(FLD_ID_SCAN, []))
        scan_status = f"✅ {len(scans)}张" if scans else "⚠️  无"
        print(f"{i:<4} {rec['record_id']:<22} {name:<20} {email:<32} {scan_status}")


# ── 主逻辑 ────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="合同生成脚本")
    parser.add_argument("--name",      help="按姓名查找（模糊匹配）")
    parser.add_argument("--record-id", dest="record_id", help="按 record_id 精确查找")
    parser.add_argument("--list",      action="store_true", help="列出所有记录")
    parser.add_argument("--dry-run",   action="store_true", help="只打印变量，不生成文件")
    parser.add_argument("--send",      action="store_true", help="生成后发送邮件")
    parser.add_argument("--draft",     action="store_true", help="生成后保存草稿，VM 双击 .eml 后点发送")
    parser.add_argument("--yes",       action="store_true", help="跳过交互确认")
    args = parser.parse_args()

    print("拉取合同信息收集表...")
    collect_records = fetch_collect_records()
    print(f"  共 {len(collect_records)} 条")

    if args.list:
        list_records(collect_records)
        return

    # ── 选候选人 ──
    target = None
    if args.record_id:
        for rec in collect_records:
            if rec["record_id"] == args.record_id:
                target = rec; break
        if not target:
            print(f"❌ 未找到 record_id={args.record_id}"); sys.exit(1)
    elif args.name:
        matches = [r for r in collect_records
                   if args.name.lower() in extract_text(r["fields"].get(FLD_NAME, "")).lower()]
        if not matches:
            print(f"❌ 未找到姓名包含「{args.name}」的记录"); sys.exit(1)
        if len(matches) > 1:
            print(f"⚠️  找到 {len(matches)} 条，请用 --record-id 精确指定：")
            list_records(matches); sys.exit(1)
        target = matches[0]
    else:
        parser.print_help(); sys.exit(0)

    fields = target["fields"]
    name       = extract_text(fields.get(FLD_NAME, "")) or "未知"
    acct_name  = extract_text(fields.get(FLD_ACCT_NAME, ""))
    email_addr = extract_text(fields.get(FLD_EMAIL, ""))

    print(f"\n=== 候选人：{name}  ({target['record_id']}) ===")

    # ── 证件扫描件检查 ──
    id_scans = extract_attachments(fields.get(FLD_ID_SCAN, []))
    if id_scans:
        print(f"✅ 证件扫描件：{len(id_scans)} 张（{', '.join(a['name'] for a in id_scans)}）")
    else:
        print("⚠️  证件扫描件：收集表中无附件，将跳过插图")

    # ── 银行账户名校验 ──
    need_confirm, msg = check_name_match(name, acct_name)
    print(f"银行账户名：{msg}")
    if need_confirm and not args.yes and not args.dry_run:
        ans = input("账户名确认无误？继续？[y/N] ").strip().lower()
        if ans != "y":
            print("已取消"); sys.exit(0)

    # ── 选合同模板 ──
    print("\n拉取合同模板表...")
    template_records = fetch_template_records()

    template_rec, template_name = pick_template_for_candidate(
        template_records,
        fields,
        auto_confirm=args.yes or args.dry_run,
    )
    if not template_rec:
        print("❌ 未选择模板，退出"); sys.exit(1)

    print(f"\n已选模板：{template_name}")
    is_company = is_company_contract(template_name)

    # 解析所需变量
    vars_text = extract_text(template_rec.get("fields", {}).get(TEMPLATE_VARS_FLD, ""))
    required_vars = parse_required_vars(vars_text)
    print(f"所需变量（{len(required_vars)} 个）：{', '.join(required_vars)}")

    # VM 补充手动变量
    vm_overrides = {}
    today = datetime.now()
    vm_overrides["签署日期"] = today.strftime("%Y-%m-%d")

    needs_vm = [v for v in required_vars if v in VM_INPUT_VARS]
    if needs_vm and not args.dry_run:
        print(f"\n以下变量需要手动输入（直接回车跳过）：")
        for var in needs_vm:
            hint = VM_INPUT_VARS[var]
            val = input(f"  {var}（{hint}）: ").strip()
            if val:
                vm_overrides[var] = val

    # 构建变量替换表
    var_map, missing, filled, empty, unmatched = build_var_map(fields, required_vars, vm_overrides)

    # ── 变量填充报告 ──────────────────────────────────────────
    print("\n" + "─" * 55)
    print(f"  变量填充报告（共 {len(required_vars)} 个变量）")
    print("─" * 55)

    if filled:
        print(f"\n✅ 已成功填充（{len(filled)} 个）：")
        for var in filled:
            key = f"{{{{{var}}}}}"
            print(f"   {var}  →  {var_map[key]}")

    if empty:
        print(f"\n⚠️  字段存在但值为空，需人工二次确认（{len(empty)} 个）：")
        for var in empty:
            print(f"   {var}  →  （资源商未填写该字段）")

    if unmatched:
        print(f"\n❌ 变量名在映射表中未找到对应字段，需人工处理（{len(unmatched)} 个）：")
        for var in unmatched:
            print(f"   {var}  →  （未知变量，请检查 field_mapping.py）")

    print("─" * 55)

    if empty or unmatched:
        total_bad = len(empty) + len(unmatched)
        print(f"\n⚠️  共 {total_bad} 个变量未能自动填充，生成的合同中这些位置将保留空白。")
        print(   "   请在合同生成后人工核查以下字段：")
        for var in empty + unmatched:
            print(f"   · {var}")
        if not args.yes and not args.dry_run:
            ans = input("\n是否继续生成合同？[y/N] ").strip().lower()
            if ans != "y":
                print("已取消，请补全信息后重新运行。")
                sys.exit(0)

    if args.dry_run:
        print("\n[DRY-RUN] 不生成文件")
        log_manual_step(
            step_name="合同生成 dry-run",
            status="skipped",
            candidate_name=name,
            candidate_record_id=target["record_id"],
            input_summary=f"模板: {template_name}",
            output_summary=f"已填充 {len(filled)} 个变量; 空字段 {len(empty)}; 未匹配 {len(unmatched)}",
        )
        return

    # ── 下载模板 docx ──
    att_list = extract_attachments(template_rec.get("fields", {}).get(TEMPLATE_ATT_FLD, []))
    if not att_list:
        print("❌ 合同模板表中无 AI合同模版 附件"); sys.exit(1)

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        template_docx = tmpdir / att_list[0]["name"]

        print(f"\n下载模板：{att_list[0]['name']}")
        # 下载 Base 附件需使用 record_id + file_token
        r = subprocess.run(
            ["lark-cli", "base", "+record-download-attachment",
             "--base-token", TEMPLATE_BASE,
             "--table-id", TEMPLATE_TABLE,
             "--record-id", template_rec["record_id"],
             "--file-token", att_list[0]["file_token"],
             "--output", template_docx.name],
            capture_output=True, text=True,
            cwd=str(template_docx.parent),
        )
        if r.returncode != 0:
            local_template = (
                Path.home()
                / "Downloads"
                / "合同模板_合同模板汇总_附件"
                / "GPT_变量标注版"
                / att_list[0]["name"]
            )
            if local_template.exists():
                template_docx.write_bytes(local_template.read_bytes())
                print(f"⚠️  飞书模板下载失败，使用本地缓存：{local_template}")
            else:
                print(f"❌ 模板下载失败：{r.stderr or r.stdout}"); sys.exit(1)
        print("✅ 模板下载完成")

        # ── 填充变量 ──
        doc = Document(template_docx)
        fill_template_vars(doc, var_map)

        # ── 插入证件扫描件 ──
        if id_scans and not is_company:
            print("\n下载并插入证件扫描件...")
            img_paths = download_id_scan_images(fields, tmpdir, target["record_id"])
            if img_paths:
                insert_id_scan(doc, img_paths)
            else:
                print("  ⚠️  下载失败，跳过插图")
        elif is_company:
            print("  （公司合同，跳过证件扫描件插入）")

        # ── 保存输出 ──
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        date_str  = today.strftime("%Y-%m-%d")
        safe_name = re.sub(r'[^\w\-\u4e00-\u9fff]', '_', name)
        out_name  = f"{date_str}_{safe_name}_{template_name}"
        output_path = OUTPUT_DIR / out_name

        doc.save(output_path)
        print(f"\n✅ 合同已生成：{output_path}")

        # 残留变量检查
        doc2 = Document(output_path)
        all_text = "\n".join(p.text for p in doc2.paragraphs)
        remaining = re.findall(r'\{\{[^}]+\}\}', all_text)
        if remaining:
            print(f"\n⚠️  二次检查：合同中仍有 {len(remaining)} 处未替换变量，请人工核查：")
            for r in remaining:
                print(f"   · {r}")
            print("   提示：这些位置在合同中将显示为空白或原始占位符。")
        else:
            print("✅ 二次检查通过：所有变量已替换完毕")
            log_manual_step(
                step_name="合同 docx 生成",
                status="done",
                candidate_name=name,
                candidate_record_id=target["record_id"],
                input_summary=f"模板: {template_name}",
                output_summary=f"文件: {output_path}",
            )

        open_docx(output_path)

        # ── 发送邮件 ──
        if args.send:
            lang = "zh" if re.search(r'[\u4e00-\u9fff]', name) else "en"
            print(f"\n发送邮件（语言：{lang}）...")
            send_email(email_addr, name, output_path, lang=lang, draft=args.draft)
            if args.draft:
                print("📝 已保存本地合同邮件草稿，不更新「合同签署」字段。")
            else:
                print("✅ 合同邮件已发送；「合同签署」字段等待签回核查节点更新。")
        else:
            print(f"\n合同已打开预览，确认无误后运行：")
            print(f"  python3 scripts/generate_contract.py --name '{name}' --send")


if __name__ == "__main__":
    main()
